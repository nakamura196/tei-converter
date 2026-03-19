#!/usr/bin/env python3
"""
仕様書Word生成スクリプト

YAMLファイルから仕様書の内容を読み込み、
スタイル付きのWord文書（.docx）を生成する。

Usage:
    python scripts/generate-spec-docx.py [input.yaml] [output.docx]

    引数なしの場合:
      入力: docs/spec.yaml
      出力: docs/仕様書.docx
"""

import sys
import yaml
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# Styling helpers
# ============================================================

def set_east_asian_font(run, font_name):
    """runに東アジアフォント（游明朝等）を設定する"""
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)


def add_run(paragraph, text, size=Pt(10.5), bold=False, font_name='Yu Mincho'):
    """スタイル付きのrunを追加する"""
    run = paragraph.add_run(text)
    run.font.size = size
    run.font.bold = bold
    run.font.name = font_name
    set_east_asian_font(run, font_name)
    return run


def remove_table_borders(table):
    """テーブルの罫線をすべて除去する"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        borders.append(border)
    tblPr.append(borders)


# ============================================================
# Document builders
# ============================================================

class SpecDocBuilder:
    """仕様書Word文書を生成するビルダー"""

    def __init__(self, config):
        self.config = config
        self.font_name = config.get('font', 'Yu Mincho')
        self.doc = Document()
        self._setup_page()
        self._setup_default_style()

    def _setup_page(self):
        page = self.config.get('page', {})
        for section in self.doc.sections:
            section.page_width = Cm(page.get('width', 21.0))
            section.page_height = Cm(page.get('height', 29.7))
            section.top_margin = Cm(page.get('margin_top', 2.5))
            section.bottom_margin = Cm(page.get('margin_bottom', 3.0))
            section.left_margin = Cm(page.get('margin_left', 2.5))
            section.right_margin = Cm(page.get('margin_right', 2.5))

    def _setup_default_style(self):
        style = self.doc.styles['Normal']
        style.font.name = self.font_name
        style.font.size = Pt(10.5)
        style.paragraph_format.line_spacing = 1.9
        rPr = style.element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), self.font_name)
        rPr.append(rFonts)

    # -- Cover page --

    def add_cover(self):
        cover = self.config.get('cover', {})
        spacing_lines = cover.get('spacing_lines', 8)

        for _ in range(spacing_lines):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.0

        # Title
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(10)
        add_run(p, cover.get('title', ''), size=Pt(22), bold=True, font_name=self.font_name)

        # Subtitle
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(140)
        add_run(p, cover.get('subtitle', '仕様書'), size=Pt(18), bold=True, font_name=self.font_name)

        # Organization
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, cover.get('organization', ''), size=Pt(14), font_name=self.font_name)

        self.doc.add_page_break()

    # -- Content elements --

    def add_doc_header(self, title, subtitle):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        add_run(p, title, size=Pt(16), bold=True, font_name=self.font_name)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(28)
        add_run(p, subtitle, size=Pt(12), bold=True, font_name=self.font_name)

    def add_section_title(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(28)
        p.paragraph_format.space_after = Pt(10)
        add_run(p, text, size=Pt(12), bold=True, font_name=self.font_name)

    def add_subsection_title(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(4)
        add_run(p, text, size=Pt(10.5), bold=True, font_name=self.font_name)

    def add_body(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run(p, text, font_name=self.font_name)

    def add_numbered_items(self, items, start=1):
        """番号付き項目をボーダーなしテーブルで追加する"""
        table = self.doc.add_table(rows=len(items), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        remove_table_borders(table)

        for i, text in enumerate(items):
            row = table.rows[i]
            # Number cell
            cell0 = row.cells[0]
            cell0.width = Cm(1.5)
            p = cell0.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            add_run(p, f'{start + i}.', font_name=self.font_name)
            # Text cell
            cell1 = row.cells[1]
            cell1.width = Cm(14.5)
            p = cell1.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            add_run(p, text, font_name=self.font_name)

    # -- Build from config --

    def build(self):
        cover = self.config.get('cover', {})
        self.add_cover()
        self.add_doc_header(cover.get('title', ''), cover.get('subtitle', '仕様書'))

        # Sections
        for section in self.config.get('sections', []):
            self.add_section_title(section['title'])

            if 'body' in section:
                self.add_body(section['body'])

            for group in section.get('groups', []):
                self.add_subsection_title(group['title'])
                start = group.get('start', 1)
                self.add_numbered_items(group['items'], start=start)

        return self.doc

    def save(self, path):
        self.doc.save(path)


# ============================================================
# Main
# ============================================================

def main():
    input_path = Path(sys.argv[1]) if len(sys.argv) > 1 else Path('docs/spec.yaml')
    output_path = Path(sys.argv[2]) if len(sys.argv) > 2 else Path('docs/仕様書.docx')

    if not input_path.exists():
        print(f'Error: {input_path} not found')
        sys.exit(1)

    with open(input_path, 'r', encoding='utf-8') as f:
        config = yaml.safe_load(f)

    builder = SpecDocBuilder(config)
    builder.build()
    builder.save(str(output_path))
    print(f'Generated: {output_path}')


if __name__ == '__main__':
    main()
